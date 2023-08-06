# import csv
import datetime

# import difflib
# import os
# import pickle
# import re
# import sys
# import typing
# from collections import Counter
import math
from typing import List, Dict, Union, Optional, Tuple
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import timedelta, date

#
# from dateutil import parser
import numpy as np
from datamaps.api import project_data_from_master
import platform
from pathlib import Path

# from dateutil.parser import ParserError
from docx import Document, table
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH

# from docx.oxml import parse_xml
# from docx.oxml.ns import nsdecls
from docx.shared import Pt, Cm, RGBColor, Inches

# from matplotlib import cm
# from matplotlib.patches import Wedge, Rectangle, Circle
from openpyxl import load_workbook, Workbook

# from openpyxl.chart import BubbleChart, Reference, Series
from openpyxl.formatting import Rule
from openpyxl.styles import Font, PatternFill, Border
from openpyxl.styles.differential import DifferentialStyle
from openpyxl.workbook import workbook
from textwrap import wrap
import logging
from pdf2image import convert_from_path
from analysis_engine.data import (
    convert_bc_stage_text,
    plus_minus_days,
    concatenate_dates,
    convert_rag_text,
    rag_txt_list,
    black_text,
    fill_colour_list,
    get_group,
    COLOUR_DICT,
    make_file_friendly,
    wd_heading,
    key_contacts,
    dca_table,
    dca_narratives,
    open_word_doc,
    set_col_widths,
    make_columns_bold,
    change_text_size,
    get_iter_list,
    dandelion_number_text,
    cal_group_angle,
    get_correct_p_data,
)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s: %(levelname)s - %(message)s",
    datefmt="%d-%b-%y %H:%M:%S",
)
logger = logging.getLogger(__name__)


class ProjectNameError(Exception):
    pass


class ProjectGroupError(Exception):
    pass


class ProjectStageError(Exception):
    pass


def _cdg_platform_docs_dir() -> Path:
    #  Cross plaform file path handling
    if platform.system() == "Linux":
        return Path.home() / "Documents" / "data_bridge"
    if platform.system() == "Darwin":
        return Path.home() / "Documents" / "data_bridge"
    else:
        return Path.home() / "Documents" / "data_bridge"


cdg_root_path = _cdg_platform_docs_dir()


def cdg_get_master_data() -> List[
    Dict[str, Union[str, int, datetime.date, float]]
]:  # how specify a list of dictionaries?
    """Returns a list of dictionaries each containing quarter data"""
    master_data_list = [
        project_data_from_master(
            cdg_root_path / "core_data/cdg_master_4_2020.xlsx", 4, 2020
        ),
        project_data_from_master(
            cdg_root_path / "core_data/cdg_master_3_2020.xlsx", 3, 2020
        ),
    ]
    return master_data_list


def cdg_get_project_information() -> Dict[str, Union[str, int]]:
    """Returns dictionary containing all project meta data"""
    return project_data_from_master(
        cdg_root_path / "core_data/cdg_project_info.xlsx", 2, 2020
    )


def place_data_into_new_master_format(master_data: Dict):  # throw away
    wb = load_workbook(cdg_root_path / "core_data/CDG_portfolio_report.xlsx")
    ws = wb.active

    for i, p in enumerate(master_data.projects):
        ws.cell(row=3, column=i + 5).value = p
        for row_num in range(2, ws.max_row + 1):
            key = ws.cell(row=row_num, column=2).value
            try:
                ws.cell(row=row_num, column=i + 5).value = master_data.data[p][key]
            except KeyError:
                pass

    return wb


BASELINE_TYPES = {
    "Re-baseline this quarter": "quarter",
}
# CDG_GROUP_DICT = {"Corporate Finance": "CF", "Group Finance": "GF"}
CDG_DIR = ["CFPD", "GF", "Digital", "SCS"]
DFT_STAGE = ["pre-SOBC", "SOBC", "OBC", "FBC"]
BC_STAGE_DICT = {
    "Strategic Outline Case": "SOBC",
    "SOBC": "SOBC",
    "pre-Strategic Outline Case": "pre-SOBC",
    "pre-SOBC": "pre-SOBC",
    "Outline Business Case": "OBC",
    "OBC": "OBC",
    "Full Business Case": "FBC",
    "FBC": "FBC",
    # older returns that require cleaning
    "Pre - SOBC": "pre-SOBC",
    "Pre Strategic Outline Business Case": "pre_SOBC",
    None: None,
    "Other": "Other",
    "Other ": "Other",
    "To be confirmed": None,
    "To be confirmed ": None,
}
DCG_DATE = datetime.date(
    2021, 2, 22
)  # ipdc date. Python date format is Year, Month, day


class CDGMaster:
    def __init__(
        self,
        master_data: List[Dict[str, Union[str, int, datetime.date, float]]],
        project_information: Dict[str, Union[str, int]],
    ) -> None:
        self.master_data = master_data
        self.project_information = project_information
        self.current_quarter = self.master_data[0].quarter
        self.current_projects = self.master_data[0].projects
        self.abbreviations = {}
        self.full_names = {}
        self.bl_info = {}
        self.bl_index = {}
        self.dft_groups = {}
        self.project_group = {}
        self.project_stage = {}
        self.quarter_list = []
        self.get_quarter_list()
        # self.get_baseline_data()
        self.check_project_information()
        self.get_project_abbreviations()
        # self.check_baselines()
        self.get_project_groups()

    def get_project_abbreviations(self) -> None:
        """gets the abbreviations for all current projects.
        held in the project info document"""
        abb_dict = {}
        fn_dict = {}
        error_case = []
        for p in self.project_information.projects:
            abb = self.project_information[p]["Abbreviations"]
            abb_dict[p] = {"abb": abb, "full name": p}
            fn_dict[abb] = p
            if abb is None:
                error_case.append(p)

        if error_case:
            for p in error_case:
                logger.critical("No abbreviation provided for " + p + ".")
            raise ProjectNameError(
                "Abbreviations must be provided for all projects in project_info. Program stopping. Please amend"
            )

        self.abbreviations = abb_dict
        self.full_names = fn_dict

    def get_baseline_data(self) -> None:
        """
        Returns the two dictionaries baseline_info and baseline_index for all projects for all
        baseline types
        """

        baseline_info = {}
        baseline_index = {}

        for b_type in list(BASELINE_TYPES.keys()):
            project_baseline_info = {}
            project_baseline_index = {}
            for name in self.current_projects:
                bc_list = []
                lower_list = []
                for i, master in reversed(list(enumerate(self.master_data))):
                    if name in master.projects:
                        try:
                            approved_bc = master.data[name][b_type]
                            quarter = str(master.quarter)
                        # exception handling in here in case data keys across masters are not consistent.
                        except KeyError:
                            print(
                                str(b_type)
                                + " keys not present in "
                                + str(master.quarter)
                            )
                        if approved_bc == "YES":
                            bc_list.append(approved_bc)
                            lower_list.append((approved_bc, quarter, i))
                    else:
                        pass
                for i in reversed(range(2)):
                    try:
                        # if name in self.master_data[i].projects:
                        approved_bc = self.master_data[i][name][b_type]
                        quarter = str(self.master_data[i].quarter)
                        lower_list.append((approved_bc, quarter, i))
                    # TODO tidy this
                    except IndexError:
                        # else:
                        #     quarter = str(self.master_data[i].quarter)
                        lower_list.append((None, "LAST", None))

                index_list = []
                for x in lower_list:
                    index_list.append(x[2])

                project_baseline_info[name] = list(reversed(lower_list))
                project_baseline_index[name] = list(reversed(index_list))

            baseline_info[BASELINE_TYPES[b_type]] = project_baseline_info
            baseline_index[BASELINE_TYPES[b_type]] = project_baseline_index

        self.bl_info = baseline_info
        self.bl_index = baseline_index

    def check_project_information(self) -> None:
        """Checks that project names in master are present/the same as in project info.
        Stops the programme if not"""
        error_cases = []
        for p in self.current_projects:
            if p not in self.project_information.projects:
                error_cases.append(p)

        if error_cases:
            for p in error_cases:
                logger.critical(p + " has not been found in the project_info document.")
            raise ProjectNameError(
                "Project names in "
                + str(self.master_data[0].quarter)
                + " master and project_info must match. Program stopping. Please amend."
            )
        else:
            logger.info("The latest master and project information match")

    def check_baselines(self) -> None:
        """checks that projects have the correct baseline information. stops the
        programme if baselines are missing"""
        # work through best way to stop the programme.
        for v in BASELINE_TYPES.values():
            for p in self.current_projects:
                baselines = self.bl_index[v][p]
                if len(baselines) <= 2:
                    print(
                        p
                        + " does not have a baseline point for "
                        + v
                        + " this could cause the programme to "
                        "crash. Therefore the programme is stopping. "
                        "Please amend the data for " + p + " so that "
                        " it has at least one baseline point for " + v
                    )
            else:
                continue
            break

    def get_project_groups(self) -> None:
        """gets the groups that projects are part of e.g. business case
        stage or dft group"""

        raw_dict = {}
        raw_list = []
        group_list = []
        stage_list = []
        for i, master in enumerate(self.master_data):
            lower_dict = {}
            for p in master.projects:
                dft_group = self.project_information[p]["Directorate"]  # different groups cleaned here
                if dft_group is None:
                    logger.critical(
                        str(p) + " does not have a Group value in the project information document."
                    )
                    raise ProjectGroupError(
                        "Program stopping as this could cause a crash. Please check project Group info."
                    )
                if dft_group not in CDG_DIR:
                    logger.critical(
                        str(p) + " Group value is " + str(dft_group) + " . This is not a recognised group"
                    )
                    raise ProjectGroupError(
                        "Program stopping as this could cause a crash. Please check project Group info."
                    )
                stage = BC_STAGE_DICT[master[p]["CDG approval point"]]
                raw_list.append(("group", dft_group))
                raw_list.append(("stage", stage))
                lower_dict[p] = dict(raw_list)
                group_list.append(dft_group)
                stage_list.append(stage)

            raw_dict[str(master.quarter)] = lower_dict

        group_list = list(set(group_list))
        stage_list = list(set(stage_list))

        group_dict = {}
        for i, quarter in enumerate(raw_dict.keys()):
            lower_g_dict = {}
            for group_type in group_list:
                g_list = []
                for p in raw_dict[quarter].keys():
                    p_group = raw_dict[quarter][p]["group"]
                    if p_group == group_type:
                        g_list.append(p)
                lower_g_dict[group_type] = g_list

            # Removed as Tier projects unlikely to be on gmpp
            # gmpp_list = []
            # for p in self.master_data[i].projects:
            #     gmpp = self.master_data[i].data[p]["GMPP - IPA ID Number"]
            #     if gmpp is not None:
            #         gmpp_list.append(p)
            #     lower_g_dict["GMPP"] = gmpp_list

            group_dict[quarter] = lower_g_dict

        stage_dict = {}
        for quarter in raw_dict.keys():
            lower_s_dict = {}
            for stage_type in stage_list:
                s_list = []
                for p in raw_dict[quarter].keys():
                    p_stage = raw_dict[quarter][p]["stage"]
                    if p_stage == stage_type:
                        s_list.append(p)
                if stage_type is None:
                    if s_list:
                        if quarter == self.current_quarter:
                            for x in s_list:
                                logger.critical(str(x) + " has no IPDC stage date")
                                raise ProjectStageError("Programme stopping as this could cause incomplete analysis")
                        else:
                            for x in s_list:
                                logger.warning(
                                    "In " + str(quarter) + " master " + str(x)
                                    + " IPDC stage data is currently None. Please amend."
                                )
                lower_s_dict[stage_type] = s_list
            stage_dict[quarter] = lower_s_dict

        self.dft_groups = group_dict
        self.project_stage = stage_dict

    def get_quarter_list(self) -> None:
        output_list = []
        for master in self.master_data:
            output_list.append(str(master.quarter))
        self.quarter_list = output_list


class CDGCostData:
    def __init__(self, master: CDGMaster, **kwargs):
        self.master = master
        self.baseline_type = "ipdc_costs"
        self.kwargs = kwargs
        self.start_group = []
        self.group = []
        self.iter_list = []
        self.c_totals = {}
        self.wlc_dict = {}
        self.wlc_change = {}
        # self.stack_p = {}
        self.get_cost_totals()
        # self.get_wlc_data()
        # self.get_stackplot_data()

    def get_cost_totals(self) -> None:

        self.iter_list = get_iter_list(self.kwargs, self.master)
        self.start_group = get_group(self.master, self.iter_list[0], self.kwargs)
        lower_dict = {}

        for tp in self.iter_list:
            self.group = get_group(self.master, tp, self.kwargs)
            profiled = 0
            for project_name in self.group:
                p_data = get_correct_p_data(
                    self.kwargs, self.master, self.baseline_type, project_name, tp
                )
                try:
                    profiled += p_data["Total Forecast"]
                except TypeError:
                    pass

            lower_dict[tp] = {
                "total": profiled,
            }

        self.c_totals = lower_dict

    def get_wlc_data(self) -> None:
        """central point in code which
        calculates the quarters total
        filters projects by group in order of size wlc"""
        self.iter_list = get_iter_list(self.kwargs, self.master)
        wlc_dict = {}
        for tp in self.iter_list:
            #  for need groups of groups.  Not consistent with steps for
            #  other functions in this class. currently only in use for dandelion
            if "group" in self.kwargs:
                self.group = self.kwargs["group"]
            elif "stage" in self.kwargs:
                self.group = self.kwargs["stage"]
            wlc_dict = {}
            p_total = 0  # portfolio total

            for i, g in enumerate(self.group):
                l_group = get_group(self.master, tp, self.kwargs, i)  # lower group
                g_total = 0
                l_g_l = []  # lower group list
                for p in l_group:
                    p_data = get_correct_p_data(
                        self.kwargs, self.master, self.baseline_type, p, tp
                    )
                    wlc = p_data["Total Forecast"]
                    if isinstance(wlc, (float, int)) and wlc is not None and wlc != 0:
                        if wlc > 50000:
                            logger.info(
                                tp
                                + ", "
                                + str(p)
                                + " is £"
                                + str(round(wlc))
                                + " please check this is correct. For now analysis_engine has recorded it as £0"
                            )
                        # wlc_dict[p] = wlc
                    if wlc == 0:
                        logger.info(
                            tp
                            + ", "
                            + str(p)
                            + " wlc is currently £"
                            + str(wlc)
                            + " note this is key information that should be provided by the project"
                        )
                        # wlc_dict[p] = wlc
                    if wlc is None:
                        logger.info(
                            tp
                            + ", "
                            + str(p)
                            + " wlc is currently None note this is key information that should be provided by the project"
                        )
                        wlc = 0

                    l_g_l.append((wlc, p))
                    g_total += wlc

                wlc_dict[g] = list(reversed(sorted(l_g_l)))
                p_total += g_total

            wlc_dict["total"] = p_total
            wlc_dict[tp] = wlc_dict

        self.wlc_dict = wlc_dict

    def calculate_wlc_change(self) -> None:
        wlc_change_dict = {}
        for i, tp in enumerate(self.wlc_dict.keys()):
            p_wlc_change_dict = {}
            for p in self.wlc_dict[tp].keys():
                wlc_one = self.wlc_dict[tp][p]
                try:
                    wlc_two = self.wlc_dict[self.iter_list[i + 1]][p]
                    try:
                        percentage_change = int(((wlc_one - wlc_two) / wlc_one) * 100)
                        p_wlc_change_dict[p] = percentage_change
                    except ZeroDivisionError:
                        logger.info(
                            "As "
                            + str(p)
                            + " has no wlc total figure for "
                            + tp
                            + " change has been calculated as zero"
                        )
                except IndexError:  # handles NoneTypes.
                    pass

            wlc_change_dict[tp] = p_wlc_change_dict

        self.wlc_change = wlc_change_dict


def cdg_overall_dashboard(master: CDGMaster, wb: Workbook) -> Workbook:
    wb = load_workbook(wb)
    ws = wb.worksheets[0]

    for row_num in range(2, ws.max_row + 1):
        project_name = ws.cell(row=row_num, column=2).value
        if project_name in master.current_projects:
            """BC Stage"""
            bc_stage = master.master_data[0].data[project_name]["CDG approval point"]
            # ws.cell(row=row_num, column=4).value = convert_bc_stage_text(bc_stage)
            ws.cell(row=row_num, column=3).value = convert_bc_stage_text(bc_stage)
            try:
                bc_stage_lst_qrt = master.master_data[1].data[project_name][
                    "CDG approval point"
                ]
                if bc_stage != bc_stage_lst_qrt:
                    # ws.cell(row=row_num, column=4).font = Font(
                    #     name="Arial", size=10, color="00fc2525"
                    # )
                    ws.cell(row=row_num, column=3).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )
            except (KeyError, IndexError):
                pass

            """planning stage"""
            plan_stage = master.master_data[0].data[project_name]["Project stage"]
            # ws.cell(row=row_num, column=5).value = plan_stage
            ws.cell(row=row_num, column=4).value = plan_stage
            try:
                plan_stage_lst_qrt = master.master_data[1].data[project_name][
                    "Project stage"
                ]
                if plan_stage != plan_stage_lst_qrt:
                    # ws.cell(row=row_num, column=5).font = Font(
                    #     name="Arial", size=10, color="00fc2525"
                    # )
                    ws.cell(row=row_num, column=4).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )
            except (KeyError, IndexError):
                pass

            """Total WLC"""
            wlc_now = master.master_data[0].data[project_name]["Total Forecast"]
            # ws.cell(row=row_num, column=6).value = wlc_now
            ws.cell(row=row_num, column=5).value = wlc_now
            """WLC variance against lst quarter"""
            try:
                wlc_lst_quarter = master.master_data[1].data[project_name][
                    "Total Forecast"
                ]
                diff_lst_qrt = wlc_now - wlc_lst_quarter
                if float(diff_lst_qrt) > 0.49 or float(diff_lst_qrt) < -0.49:
                    # ws.cell(row=row_num, column=7).value = diff_lst_qrt
                    ws.cell(row=row_num, column=6).value = diff_lst_qrt
                else:
                    # ws.cell(row=row_num, column=7).value = "-"
                    ws.cell(row=row_num, column=6).value = "-"

                try:
                    percentage_change = ((wlc_now - wlc_lst_quarter) / wlc_now) * 100
                    if percentage_change > 5 or percentage_change < -5:
                        # ws.cell(row=row_num, column=7).font = Font(
                        #     name="Arial", size=10, color="00fc2525"
                        # )
                        ws.cell(row=row_num, column=6).font = Font(
                            name="Arial", size=10, color="00fc2525"
                        )
                except ZeroDivisionError:
                    pass

            except (KeyError, IndexError):
                ws.cell(row=row_num, column=6).value = "-"

            """WLC variance against baseline quarter"""
            bl = master.bl_index["quarter"][project_name][2]
            wlc_baseline = master.master_data[bl].data[project_name]["Total Forecast"]
            try:
                diff_bl = wlc_now - wlc_baseline
                if float(diff_bl) > 0.49 or float(diff_bl) < -0.49:
                    # ws.cell(row=row_num, column=8).value = diff_bl
                    ws.cell(row=row_num, column=7).value = diff_bl
                else:
                    # ws.cell(row=row_num, column=8).value = "-"
                    ws.cell(row=row_num, column=7).value = "-"
            except TypeError:  # exception is here as some projects e.g. Hs2 phase 2b have (real) written into historical totals
                pass

            try:
                percentage_change = ((wlc_now - wlc_baseline) / wlc_now) * 100
                if percentage_change > 5 or percentage_change < -5:
                    # ws.cell(row=row_num, column=8).font = Font(
                    #     name="Arial", size=10, color="00fc2525"
                    # )
                    ws.cell(row=row_num, column=7).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )

            except (
                ZeroDivisionError,
                TypeError,
            ):  # zerodivision error obvious, type error handling as above
                pass

            """vfm category now"""
            vfm_cat = master.master_data[0].data[project_name][
                "VfM Category single entry"
            ]
            # if (
            #     master.master_data[0].data[project_name]["VfM Category single entry"]
            #     is None
            # ):
            #     vfm_cat = (
            #         str(
            #             master.master_data[0].data[project_name][
            #                 "VfM Category lower range"
            #             ]
            #         )
            #         + " - "
            #         + str(
            #             master.master_data[0].data[project_name][
            #                 "VfM Category upper range"
            #             ]
            #         )
            #     )
            #     # ws.cell(row=row_num, column=10).value = vfm_cat
            #     ws.cell(row=row_num, column=8).value = vfm_cat
            #
            # else:
            #     vfm_cat = master.master_data[0].data[project_name][
            #         "VfM Category single entry"
            #     ]
            #     # ws.cell(row=row_num, column=10).value = vfm_cat
            ws.cell(row=row_num, column=8).value = vfm_cat

            """vfm category baseline"""
            bl_i = master.bl_index["quarter"][project_name][2]
            vfm_cat_baseline = master.master_data[bl_i].data[project_name][
                "VfM Category single entry"
            ]
            # try:
            #     if (
            #         master.master_data[bl_i].data[project_name][
            #             "VfM Category single entry"
            #         ]
            #         is None
            #     ):
            #         vfm_cat_baseline = (
            #             str(
            #                 master.master_data[bl_i].data[project_name][
            #                     "VfM Category lower range"
            #                 ]
            #             )
            #             + " - "
            #             + str(
            #                 master.master_data[bl_i].data[project_name][
            #                     "VfM Category upper range"
            #                 ]
            #             )
            #         )
            #         # ws.cell(row=row_num, column=11).value = vfm_cat_baseline
            #     else:
            #         vfm_cat_baseline = master.master_data[bl_i].data[project_name][
            #             "VfM Category single entry"
            #         ]
            #         # ws.cell(row=row_num, column=11).value = vfm_cat_baseline

            # except KeyError:
            #     try:
            #         vfm_cat_baseline = master.master_data[bl_i].data[project_name][
            #             "VfM Category single entry"
            #         ]
            #         # ws.cell(row=row_num, column=11).value = vfm_cat_baseline
            #     except KeyError:
            #         vfm_cat_baseline = master.master_data[bl_i].data[project_name][
            #             "VfM Category"
            #         ]
            #         # ws.cell(row=row_num, column=11).value = vfm_cat_baseline

            if vfm_cat != vfm_cat_baseline:
                if vfm_cat_baseline is None:
                    pass
                else:
                    ws.cell(row=row_num, column=8).font = Font(
                        name="Arial", size=8, color="00fc2525"
                    )

            current = master.master_data[0].data[project_name]["Project End Date"]
            try:
                last_quarter = master.master_data[1].data[project_name][
                    "Full Operations"
                ]
            except IndexError:
                pass
            bl = master.master_data[bl_i].data[project_name]["Project End Date"]
            #
            # abb = master.abbreviations[project_name]["abb"]
            # current = get_milestone_date(
            #     abb, milestones.milestone_dict, "current", " Full Operations"
            # )
            # last_quarter = get_milestone_date(
            #     abb, milestones.milestone_dict, "last", " Full Operations"
            # )
            # bl = get_milestone_date(
            #     abb, milestones.milestone_dict, "bl_one", " Full Operations"
            # )
            ws.cell(row=row_num, column=9).value = current
            if current is not None and current < DCG_DATE:
                ws.cell(row=row_num, column=9).value = "Completed"
            try:
                last_change = (current - last_quarter).days
                if last_change == 0:
                    ws.cell(row=row_num, column=10).value = "-"
                else:
                    ws.cell(row=row_num, column=10).value = plus_minus_days(last_change)
                if last_change is not None and last_change > 46:
                    ws.cell(row=row_num, column=10).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )
            except (TypeError, UnboundLocalError):
                pass
            try:
                bl_change = (current - bl).days
                if bl_change == 0:
                    ws.cell(row=row_num, column=11).value = "-"
                else:
                    ws.cell(row=row_num, column=11).value = plus_minus_days(bl_change)
                if bl_change is not None and bl_change > 85:
                    ws.cell(row=row_num, column=11).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )
            except TypeError:
                pass

            # last at/next at cdg information  removed
            try:
                ws.cell(row=row_num, column=12).value = concatenate_dates(
                    master.master_data[0].data[project_name]["Last date at CDG"],
                    DCG_DATE,
                )
                ws.cell(row=row_num, column=13).value = concatenate_dates(
                    master.master_data[0].data[project_name]["Next date at CDG"],
                    DCG_DATE,
                )
            except (KeyError, TypeError):
                print(
                    project_name
                    + " last at / next at ipdc data could not be calculated. Check data."
                )

            # """IPA DCA rating"""
            # ipa_dca = convert_rag_text(
            #     master.master_data[0].data[project_name]["GMPP - IPA DCA"]
            # )
            # ws.cell(row=row_num, column=15).value = ipa_dca
            # if ipa_dca == "None":
            #     ws.cell(row=row_num, column=15).value = ""

            """DCA rating - this quarter"""
            ws.cell(row=row_num, column=17).value = convert_rag_text(
                master.master_data[0].data[project_name]["Departmental DCA"]
            )
            """DCA rating - last qrt"""
            try:
                ws.cell(row=row_num, column=19).value = convert_rag_text(
                    master.master_data[1].data[project_name]["Departmental DCA"]
                )
            except (KeyError, IndexError):
                ws.cell(row=row_num, column=19).value = ""
            """DCA rating - 2 qrts ago"""
            try:
                ws.cell(row=row_num, column=20).value = convert_rag_text(
                    master.master_data[2].data[project_name]["Departmental DCA"]
                )
            except (KeyError, IndexError):
                ws.cell(row=row_num, column=20).value = ""
            """DCA rating - 3 qrts ago"""
            try:
                ws.cell(row=row_num, column=21).value = convert_rag_text(
                    master.master_data[3].data[project_name]["Departmental DCA"]
                )
            except (KeyError, IndexError):
                ws.cell(row=row_num, column=21).value = ""
            """DCA rating - baseline"""
            bl_i = master.bl_index["quarter"][project_name][2]
            ws.cell(row=row_num, column=23).value = convert_rag_text(
                master.master_data[bl_i].data[project_name]["Departmental DCA"]
            )

        """list of columns with conditional formatting"""
        list_columns = ["o", "q", "s", "t", "u", "w"]

        """same loop but the text is black. In addition these two loops go through the list_columns list above"""
        for column in list_columns:
            for i, dca in enumerate(rag_txt_list):
                text = black_text
                fill = fill_colour_list[i]
                dxf = DifferentialStyle(font=text, fill=fill)
                rule = Rule(
                    type="containsText", operator="containsText", text=dca, dxf=dxf
                )
                for_rule_formula = 'NOT(ISERROR(SEARCH("' + dca + '",' + column + "5)))"
                rule.formula = [for_rule_formula]
                ws.conditional_formatting.add(column + "5:" + column + "60", rule)

        for row_num in range(2, ws.max_row + 1):
            for col_num in range(5, ws.max_column + 1):
                if ws.cell(row=row_num, column=col_num).value == 0:
                    ws.cell(row=row_num, column=col_num).value = "-"

    return wb


# class DandelionDataOLD:
#     def __init__(self, master: Master, **kwargs):
#         self.master = master
#         self.kwargs = kwargs
#         self.baseline_type = "ipdc_costs"
#         self.group = []
#         self.iter_list = []
#         self.d_data = {}
#         self.d_list = []
#         self.get_data()
#
#     def get_data(self):
#         self.group = self.master.master_data[0].projects
#         input_g_list = ["CF", "GF"]  # first outer circle
#         # input_g_list = ["FBC", "OBC", "SOBC", "pre-SOBC"]  # first outer circle
#         # cal group angle
#         if len(input_g_list) == 2:
#             g_ang_list = [0, 225, 315]
#         # g_ang = 270/len(input_g_list)  # group angle
#         # g_ang_list = []
#         # for i in range(6):
#         #     g_ang_list.append(g_ang * i)
#         # del g_ang_list[4]
#
#         dft_g_list = []
#         dft_g_dict = {}
#         dft_l_group_dict = {}
#         p_total = 0  # portfolio total
#         for i, g in enumerate(input_g_list):
#             dft_l_group = self.master.dft_groups["Q3 20/21"][g]
#             g_total = 0
#             dft_l_group_list = []
#             for p in dft_l_group:
#                 p_data = self.master.master_data[0].data[p]
#                 b_size = p_data["Total Forecast"]
#                 if b_size is None:
#                     b_size = 25
#                 rag = p_data["Departmental DCA"]
#                 colour = COLOUR_DICT[convert_rag_text(rag)]
#                 g_total += b_size
#                 dft_l_group_list.append((math.sqrt(b_size), colour, self.master.abbreviations[p]['abb']))
#             # group data
#             x_axis = 0 + 120 * math.cos(math.radians(g_ang_list[i + 1]))
#             y_axis = 0 + 100 * math.sin(math.radians(g_ang_list[i + 1]))
#             # list is tuple axis point, bubble size, colour, line style, line color, text position
#             dft_g_list.append(((x_axis, y_axis),
#                                math.sqrt(g_total),
#                                '#a0c1d5',
#                                g,
#                                'dashed',
#                                'grey',
#                                ('center', 'center')))
#             dft_g_dict[g] = [(x_axis, y_axis), math.sqrt(g_total)/3]  # used for placement of circles
#             # project data
#             dft_l_group_dict[g] = list(reversed(sorted(dft_l_group_list)))
#             #portfolio data
#             p_total += g_total
#         dft_g_list.append(((0, 0), math.sqrt(p_total), '#a0c1d5', "CDG Portfolio", "dashed", "grey", ('center', 'center')))
#
#         for g in dft_l_group_dict.keys():
#             lg = dft_l_group_dict[g]  # local group
#             ang = 360/len(lg)
#             ang_list = []
#             for i in range(len(lg)+1):
#                 ang_list.append(ang*i)
#             for i, p in enumerate(lg):
#                 a = dft_g_dict[g][0][0]
#                 b = dft_g_dict[g][0][1]
#                 if len(lg) <= 8:
#                     x_axis = a + (dft_g_dict[g][1] + 20) * math.cos(math.radians(ang_list[i+1]))
#                     y_axis = b + (dft_g_dict[g][1] + 20) * math.sin(math.radians(ang_list[i+1]))
#                 else:
#                     x_axis = a + (dft_g_dict[g][1] + 40) * math.cos(math.radians(ang_list[i + 1]))
#                     y_axis = b + (dft_g_dict[g][1] + 40) * math.sin(math.radians(ang_list[i + 1]))
#                 b_size = p[0]
#                 colour = p[1]
#                 name = p[2]
#                 if 280 >= ang_list[i+1] >= 80:
#                     text_angle = ('right', 'bottom')
#                 if 100 >= ang_list[i+1] or ang_list[i+1] >= 260:
#                     text_angle = ('left', 'bottom')
#                 if 279 >= ang_list[i+1] >= 261:
#                     text_angle = ('center', 'top')
#                 if 99 >= ang_list[i+1] >= 81:
#                     text_angle = ('center', 'bottom')
#                 dft_g_list.append(((x_axis, y_axis), b_size, colour, name, "solid", colour, text_angle))
#
#         self.d_list = dft_g_list


def cdg_get_dandelion_type_total(
        master: CDGMaster, tp: str, g: str or List[str], kwargs
) -> int or str:  # Note no **kwargs as existing kwargs dict passed in
    if "type" in kwargs:
        if kwargs["type"] == "remaining":
            cost = CDGCostData(master, quarter=[tp], group=[g])  # group costs data
            return cost.c_totals[tp]["prof"] + cost.c_totals[tp]["unprof"]
        if kwargs["type"] == "spent":
            cost = CDGCostData(master, quarter=[tp], group=[g])  # group costs data
            return cost.c_totals[tp]["spent"]
        # No benefits yet
        # if kwargs["type"] == "benefits":
        #     benefits = BenefitsData(master, quarter=[tp], group=[g])
        #     return benefits.b_totals[tp]["total"]

    else:
        cost = CDGCostData(master, **kwargs)  # group costs data
        return cost.c_totals[tp]["total"]


class CDGDandelionData:
    def __init__(self, master: CDGMaster, **kwargs):
        self.master = master
        self.kwargs = kwargs
        self.baseline_type = "ipdc_costs"
        self.group = []
        self.iter_list = []
        self.d_data = {}
        self.get_data()

    def get_data(self):
        self.iter_list = get_iter_list(self.kwargs, self.master)
        for tp in self.iter_list:
            #  for dandelion need groups of groups.
            if "group" in self.kwargs:
                self.group = self.kwargs["group"]
            elif "stage" in self.kwargs:
                self.group = self.kwargs["stage"]

            if len(self.group) == 5:
                g_ang_l = [260, 310, 360, 50, 100]  # group angle list
            if len(self.group) == 4:
                g_ang_l = [260, 326, 32, 100]
            if len(self.group) == 3:
                g_ang_l = [280, 360, 80]
            if len(self.group) == 2:
                g_ang_l = [290, 70]
            if len(self.group) == 1:
                pass
            g_d = {}  # group dictionary. first outer circle.
            l_g_d = {}  # lower group dictionary

            pf_wlc = cdg_get_dandelion_type_total(
                self.master, tp, self.group, self.kwargs
            )  # portfolio wlc
            if "pc" in self.kwargs:  # pc portfolio colour
                pf_colour = COLOUR_DICT[self.kwargs["pc"]]
                pf_colour_edge = COLOUR_DICT[self.kwargs["pc"]]
            else:
                pf_colour = "#FFFFFF" # white
                pf_colour_edge = "grey"
            pf_text = "Portfolio\n" + dandelion_number_text(
                pf_wlc
            )  # option to specify pf name

            ## center circle
            g_d["portfolio"] = {
                "axis": (0, 0),
                "r": math.sqrt(pf_wlc),
                "colour": pf_colour,
                "text": pf_text,
                "fill": "solid",
                "ec": pf_colour_edge,
                "alignment": ("center", "center"),
            }

            ## first outer circle
            for i, g in enumerate(self.group):
                self.kwargs["group"] = [g]
                g_wlc = cdg_get_dandelion_type_total(self.master, tp, g, self.kwargs)
                if len(self.group) > 1:
                    y_axis = 0 + (
                        (math.sqrt(pf_wlc) * 3.25) * math.sin(math.radians(g_ang_l[i]))
                    )
                    x_axis = 0 + (math.sqrt(pf_wlc) * 2.75) * math.cos(
                        math.radians(g_ang_l[i])
                    )
                    g_text = g + "\n" + dandelion_number_text(g_wlc)  # group text
                    if g_wlc == 0:
                        g_wlc = pf_wlc/20
                    g_d[g] = {
                        "axis": (y_axis, x_axis),
                        "r": math.sqrt(g_wlc),
                        "wlc": g_wlc,
                        "colour": "#FFFFFF",
                        "text": g_text,
                        "fill": "dashed",
                        "ec": "grey",
                        "alignment": ("center", "center"),
                        "angle": g_ang_l[i],
                    }

                else:
                    g_d = {}
                    pf_wlc = g_wlc * 3
                    g_text = g + "\n" + dandelion_number_text(g_wlc)  # group text
                    if g_wlc == 0:
                        g_wlc = 5
                    g_d[g] = {
                        "axis": (0, 0),
                        "r": math.sqrt(g_wlc),
                        "wlc": g_wlc,
                        "colour": "#FFFFFF",
                        "text": g_text,
                        "fill": "dashed",
                        "ec": "grey",
                        "alignment": ("center", "center"),
                    }

            ## second outer circle
            for i, g in enumerate(self.group):
                self.kwargs["group"] = [g]
                group = get_group(self.master, tp, self.kwargs)  # lower group
                p_list = []
                for p in group:
                    self.kwargs["group"] = [p]
                    p_value = cdg_get_dandelion_type_total(
                        self.master, tp, p, self.kwargs
                    )  # project wlc
                    p_list.append((p_value, p))
                l_g_d[g] = list(reversed(sorted(p_list)))

            for g in self.group:
                g_wlc = g_d[g]["wlc"]
                g_radius = g_d[g]["r"]
                g_y_axis = g_d[g]["axis"][0]  # group y axis
                g_x_axis = g_d[g]["axis"][1]  # group x axis
                try:
                    p_values_list, p_list = zip(*l_g_d[g])
                except ValueError:  # handles no projects in l_g_d list
                    continue
                if len(p_list) > 2 or len(self.group) == 1:
                    ang_l = cal_group_angle(360, p_list, all=True)
                else:
                    if len(p_list) == 1:
                        ang_l = [g_d[g]["angle"]]
                    if len(p_list) == 2:
                        ang_l = [g_d[g]["angle"], g_d[g]["angle"] + 60]

                for i, p in enumerate(p_list):
                    p_value = p_values_list[i]
                    p_data = get_correct_p_data(
                        self.kwargs, self.master, self.baseline_type, p, tp
                    )
                    # change confidence type here
                    # SRO Schedule Confidence
                    # Departmental DCA
                    # SRO Benefits RAG
                    rag = p_data["Departmental DCA"]
                    colour = COLOUR_DICT[convert_rag_text(rag)]  # bubble colour
                    project_text = (
                        self.master.abbreviations[p]["abb"]
                        + "\n"
                        + dandelion_number_text(p_value)
                    )
                    if p_value == 0:
                        p_value = 2
                    # No GMPP projects
                    # if p in self.master.dft_groups[tp]["GMPP"]:
                    #     edge_colour = "#000000"  # edge of bubble
                    # else:
                    if colour == "#FFFFFF":
                        edge_colour = "grey"
                    else:
                        edge_colour = colour

                    # multi = math.sqrt(pf_wlc/g_wlc)  # multiplier
                    # multi = (1 - (g_wlc / pf_wlc)) * 3
                    try:
                        if len(p_list) >= 15:
                            multi = (pf_wlc / g_wlc) ** (1.2)
                        elif 14 >= len(p_list) >= 11:
                            multi = (pf_wlc / g_wlc) ** (1.0 / 2.0)  # square root
                        else:
                            if g_wlc/pf_wlc >= 0.33:
                                multi = (pf_wlc / g_wlc) ** (1.0 / 2.0)  # cube root
                            else:
                                multi = (pf_wlc / g_wlc) ** (1.0 / 3.0)  # cube root
                        p_y_axis = g_y_axis + (g_radius * multi) * math.sin(
                            math.radians(ang_l[i])
                        )
                        p_x_axis = g_x_axis + (g_radius * multi) * math.cos(
                            math.radians(ang_l[i])
                        )
                    except ZeroDivisionError:
                        p_y_axis = g_y_axis + 100 * math.sin(math.radians(ang_l[i]))
                        p_x_axis = g_x_axis + 100 * math.cos(math.radians(ang_l[i]))

                    if 185 >= ang_l[i] >= 175:
                        text_angle = ("center", "top")
                    if 5 >= ang_l[i] or 355 <= ang_l[i]:
                        text_angle = ("center", "bottom")
                    if 174 >= ang_l[i] >= 6:
                        text_angle = ("left", "center")
                    if 354 >= ang_l[i] >= 186:
                        text_angle = ("right", "center")

                    try:
                        t_multi = (g_wlc / p_value) ** (1.0 / 4.0)
                        # t_multi = (1 - (p_value/g_wlc)) * 2  # text multiplier
                    except ZeroDivisionError:
                        t_multi = 1
                    yx_text_position = (
                        p_y_axis
                        + (math.sqrt(p_value) * t_multi)
                        * math.sin(math.radians(ang_l[i])),
                        p_x_axis
                        + (math.sqrt(p_value) * t_multi)
                        * math.cos(math.radians(ang_l[i])),
                    )

                    g_d[p] = {
                        "axis": (p_y_axis, p_x_axis),
                        "r": math.sqrt(p_value),
                        "wlc": p_value,
                        "colour": colour,
                        "text": project_text,
                        "fill": "solid",
                        "ec": edge_colour,
                        "alignment": text_angle,
                        "tp": yx_text_position,
                    }

        self.d_data = g_d


def cdg_make_a_dandelion_auto(dl: CDGDandelionData, **kwargs):
    fig, ax = plt.subplots()
    fig.set_size_inches(18.5, 10.5)
    # ax.set_facecolor('xkcd:salmon')  # TBC if face colour is required
    # title = get_chart_title(dl_data, kwargs, "dandelion")
    # plt.suptitle(title, fontweight="bold", fontsize=10)

    for c in dl.d_data.keys():
        circle = plt.Circle(
            dl.d_data[c]["axis"],  # x, y position
            radius=dl.d_data[c]["r"],
            fc=dl.d_data[c]["colour"],  # face colour
            # linestyle=dl.d_data[c]["fill"],
            ec=dl.d_data[c]["ec"],  # edge colour
            zorder=2,
        )
        ax.add_patch(circle)
        try:
            ax.annotate(
                dl.d_data[c]["text"],  # text
                xy=dl.d_data[c]["axis"],  # x, y position
                xycoords="data",
                xytext=dl.d_data[c]["tp"],  # text position
                fontsize=7,
                # textcoords="offset pixels",
                horizontalalignment=dl.d_data[c]["alignment"][0],
                verticalalignment=dl.d_data[c]["alignment"][1],
                zorder=3,
            )
        except KeyError:
            ax.annotate(
                dl.d_data[c]["text"],  # text
                xy=dl.d_data[c]["axis"],  # x, y position
                fontsize=9,
                horizontalalignment=dl.d_data[c]["alignment"][0],
                verticalalignment=dl.d_data[c]["alignment"][1],
                weight="bold",  # bold here as will be group text
                zorder=3,
            )

    # place lines
    line_clr = "#ececec"
    line_style = "dashed"
    for i, g in enumerate(dl.group):
        dl.kwargs["group"] = [g]
        ax.arrow(
            0,
            0,
            dl.d_data[g]["axis"][0],
            dl.d_data[g]["axis"][1],
            fc=line_clr,
            ec=line_clr,
            # linestyle=line_style,
            zorder=1,
        )

        lower_g = get_group(dl.master, dl.iter_list[0], dl.kwargs)
        for p in lower_g:
            ax.arrow(
                dl.d_data[g]["axis"][0],
                dl.d_data[g]["axis"][1],
                dl.d_data[p]["axis"][0] - dl.d_data[g]["axis"][0],
                dl.d_data[p]["axis"][1] - dl.d_data[g]["axis"][1],
                fc=line_clr,
                ec=line_clr,
                # linestyle=line_style,
                zorder=1,
            )

    # ax.axes.set_xticks([])
    # ax.axes.set_yticks([])
    plt.axis("scaled")
    plt.axis("off")

    if "chart" in kwargs:
        if kwargs["chart"]:
            plt.show()

    return fig


# def convert_pdf_to_png():
#     pages = convert_from_path(root_path / "output/dandelion.pdf", 500)
#     for page in pages:
#         page.save(root_path / "output/dandelion.jpeg", "JPEG")


def cdg_compile_p_report(
    doc: Document,
    project_info: Dict[str, Union[str, int, date, float]],
    master: CDGMaster,
    project_name: str,
) -> Document:
    wd_heading(doc, project_info, project_name)
    key_contacts(doc, master, project_name)
    dca_table(doc, master, project_name)
    cdg_project_report_meta_data(doc, master, project_name)
    dca_narratives(doc, master, project_name)
    # costs = CostData(master, group=[project_name], baseline=["standard"])
    # benefits = BenefitsData(master, project_name)
    # milestones = MilestoneData(master, group=[project_name], baseline=["standard"])
    # project_report_meta_data(doc, costs, milestones, benefits, project_name)
    # change_word_doc_landscape(doc)
    # cost_profile = cost_profile_graph(costs, show="No")
    # put_matplotlib_fig_into_word(doc, cost_profile, transparent=False, size=8)
    # total_profile = total_costs_benefits_bar_chart(costs, benefits, show="No")
    # put_matplotlib_fig_into_word(doc, total_profile, transparent=False, size=8)
    # #  handling of no milestones within filtered period.
    # ab = master.abbreviations[project_name]["abb"]
    # try:
    #     # milestones.get_milestones()
    #     # milestones.get_chart_info()
    #     milestones.filter_chart_info(dates=["1/9/2020", "30/12/2022"])
    #     milestones_chart = milestone_chart(
    #         milestones,
    #         blue_line="ipdc_date",
    #         title=ab + " schedule (2021 - 22)",
    #         show="No",
    #     )
    #     put_matplotlib_fig_into_word(doc, milestones_chart, transparent=False, size=8)
    #     # print_out_project_milestones(doc, milestones, project_name)
    # except ValueError:  # extends the time period.
    #     milestones = MilestoneData(master, project_name)
    #     # milestones.get_milestones()
    #     # milestones.get_chart_info()
    #     milestones.filter_chart_info(dates=["1/9/2020", "30/12/2024"])
    #     milestones_chart = milestone_chart(
    #         milestones,
    #         blue_line="ipdc_date",
    #         title=ab + " schedule (2021 - 24)",
    #         show="No",
    #     )
    #     put_matplotlib_fig_into_word(doc, milestones_chart)
    # print_out_project_milestones(doc, milestones, project_name)
    # change_word_doc_portrait(doc)
    # project_scope_text(doc, master, project_name)
    return doc


def cdg_run_p_reports(master: CDGMaster, **kwargs) -> None:
    group = master.current_projects
    # group = get_group(master, str(master.current_quarter), kwargs)

    for p in group:
        print("Compiling summary for " + p)
        report_doc = open_word_doc(cdg_root_path / "input/summary_temp.docx")
        qrt = make_file_friendly(str(master.master_data[0].quarter))
        output = cdg_compile_p_report(report_doc, cdg_get_project_information(), master, p)
        abb = master.abbreviations[p]["abb"]
        output.save(
            cdg_root_path / "output/{}_report_{}.docx".format(abb, qrt)
        )  # add quarter here


def cdg_project_report_meta_data(
    doc: Document,
    master: CDGMaster,
    project_name: str,
):
    """Meta data table"""
    # doc.add_section(WD_SECTION_START.NEW_PAGE)
    # paragraph = doc.add_paragraph()
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # paragraph.add_run("key project MI").bold = True

    """Costs meta data"""
    # this chuck is pretty messy because the data is messy
    run = doc.add_paragraph().add_run("Key meta data")
    font = run.font
    font.bold = True
    # font.underline = True
    t = doc.add_table(rows=1, cols=4)
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = "WLC:"
    try:
        hdr_cells[1].text = (
            "£"
            + str(round(master.master_data[0].data[project_name]["Total Forecast"]))
            + "m"
        )
    except TypeError:
        hdr_cells[1].text = "TBC"
    hdr_cells[2].text = "Business Case"
    hdr_cells[3].text = str(
        master.master_data[0].data[project_name]["CDG approval point"]
    )

    row_cells = t.add_row().cells
    row_cells[0].text = "Income:"
    row_cells[1].text = ""
    row_cells[2].text = "VFM:"
    row_cells[3].text = str(
        master.master_data[0].data[project_name]["VfM Category single entry"]
    )

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(t, column_widths)
    # make column keys bold
    make_columns_bold([t.columns[0], t.columns[2]])
    change_text_size([t.columns[0], t.columns[1], t.columns[2], t.columns[3]], 10)

    return doc
